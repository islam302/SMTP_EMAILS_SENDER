import tkinter as tk
import pandas as pd
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
import os
import re
import sys
import glob
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk
import base64


class EmailSenderApp(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("Email Sender")
        self.geometry("800x700")
        self.configure(bg="white")

        # Create a Canvas for the entire window background
        self.canvas_bg = Canvas(self, bg='white', width=800, height=700)
        self.canvas_bg.pack(fill='both', expand=True)

        # Load and display background image
        self.load_background_image()

        # Style configuration
        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 10, 'bold'))
        self.style.configure('TLabel', background='white', foreground='#0078D7', font=('Segoe UI', 10))
        self.style.configure('TEntry', font=('Segoe UI', 10))

        self.word_template_path = None
        self.html_template_path = None
        self.excel_file_path = None
        self.email_provider = StringVar()

        self.email_credentials = self.load_email_credentials()
        self.selected_email = StringVar()

        self.create_widgets()

    def load_background_image(self):
        # Find and load the background image
        current_dir = os.path.dirname(sys.argv[0])  # For PyInstaller compatibility
        background_files = glob.glob(os.path.join(current_dir, "logo.*"))
        if background_files:
            background_image_path = background_files[0]
            if os.path.exists(background_image_path):
                try:
                    background_image = Image.open(background_image_path)
                    background_image = background_image.resize((800, 700), Image.LANCZOS)
                    self.background_photo = ImageTk.PhotoImage(background_image)
                    self.canvas_bg.create_image(0, 0, anchor=NW, image=self.background_photo)
                except Exception as e:
                    print(f"Error loading background image: {e}")
                    messagebox.showerror("Error", "Failed to load background image")
        else:
            print("Background image not found")

    def create_widgets(self):
        label_subject = ttk.Label(self, text="Subject:", font=("Segoe UI", 16, "bold"))
        self.canvas_bg.create_window(20, 320, anchor='w', window=label_subject)

        self.entry_subject = ttk.Entry(self, width=70, font=("Segoe UI", 12))
        self.canvas_bg.create_window(150, 320, anchor='w', window=self.entry_subject)

        btn_browse_word = ttk.Button(self, text="Choose Word Template", command=self.browse_word_template)
        self.canvas_bg.create_window(150, 360, anchor='w', window=btn_browse_word)

        btn_browse_html = ttk.Button(self, text="Choose HTML Template", command=self.browse_html_template)
        self.canvas_bg.create_window(150, 400, anchor='w', window=btn_browse_html)

        btn_browse_excel = ttk.Button(self, text="Choose Excel File", command=self.browse_excel_file)
        self.canvas_bg.create_window(150, 440, anchor='w', window=btn_browse_excel)

        label_provider = ttk.Label(self, text="Choose Email Provider:", font=("Segoe UI", 16, "bold"))
        self.canvas_bg.create_window(20, 480, anchor='w', window=label_provider)

        provider_combobox = ttk.Combobox(self, textvariable=self.email_provider, state='readonly',
                                         font=("Segoe UI", 12))
        provider_combobox['values'] = ("Gmail", "Outlook", "Webmail")
        provider_combobox.current(0)
        self.canvas_bg.create_window(150, 480, anchor='w', window=provider_combobox)

        label_email = ttk.Label(self, text="Choose Email Account:", font=("Segoe UI", 16, "bold"))
        self.canvas_bg.create_window(20, 520, anchor='w', window=label_email)

        email_combobox = ttk.Combobox(self, textvariable=self.selected_email, state='readonly',
                                      font=("Segoe UI", 12))
        email_combobox['values'] = list(self.email_credentials.keys())
        email_combobox.current(0)
        self.canvas_bg.create_window(150, 520, anchor='w', window=email_combobox)

        btn_send_emails = ttk.Button(self, text="Send Emails", command=self.send_emails)
        self.canvas_bg.create_window(150, 560, anchor='w', window=btn_send_emails)

        btn_exit = ttk.Button(self, text="Exit", command=self.destroy)
        self.canvas_bg.create_window(150, 600, anchor='w', window=btn_exit)

    def load_email_credentials(self):
        credentials = {}
        try:
            with open("credentials.txt", "r") as file:
                content = file.read()
                blocks = content.split("\n\n")
                for block in blocks:
                    lines = block.split("\n")
                    if len(lines) == 3:  # Assuming each block has email, password, and provider
                        encoded_email = lines[0]
                        encoded_password = lines[1]
                        encoded_provider = lines[2]

                        email = base64.b64decode(encoded_email).decode('utf-8')
                        password = base64.b64decode(encoded_password).decode('utf-8')
                        provider = base64.b64decode(encoded_provider).decode('utf-8')

                        credentials[email] = (password, provider)
        except Exception as e:
            print(f"Error loading email credentials: {e}")
            messagebox.showerror("Error", "Failed to load email credentials")
        return credentials

    def browse_word_template(self):
        filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if filename:
            self.word_template_path = filename
            messagebox.showinfo("Word Template Selected", f"Selected template: {filename}")

    def browse_html_template(self):
        filename = filedialog.askopenfilename(filetypes=[("HTML files", "*.html")])
        if filename:
            self.html_template_path = filename
            messagebox.showinfo("HTML Template Selected", f"Selected template: {filename}")

    def browse_excel_file(self):
        filename = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if filename:
            self.excel_file_path = filename
            messagebox.showinfo("Excel File Selected", f"Selected file: {filename}")

    # def send_emails(self):
    #     try:
    #         word_path = self.word_template_path
    #         html_path = self.html_template_path
    #         excel_path = self.excel_file_path
    #         email = self.selected_email.get()
    #         password, provider = self.email_credentials.get(email, ("", ""))  # Get password and provider tuple
    #
    #         # Validate input
    #         if not word_path or not os.path.exists(word_path):
    #             messagebox.showerror("Error", "Please choose a Word template file")
    #             return
    #
    #         if not excel_path or not os.path.exists(excel_path):
    #             messagebox.showerror("Error", "Please choose an Excel file")
    #             return
    #
    #         if not email or not password:
    #             messagebox.showerror("Error", "Please choose a valid email account")
    #             return
    #
    #         # Set SMTP server and port based on provider
    #         if provider == "Gmail":
    #             smtp_server = 'smtp.gmail.com'
    #             smtp_port = 465
    #         elif provider == "Outlook":
    #             smtp_server = 'smtp-mail.outlook.com'
    #             smtp_port = 587
    #         elif provider == "Webmail":
    #             smtp_server = 'mail.una-oic.org'
    #             smtp_port = 465
    #         else:
    #             print("Error", f"Unknown email provider: {provider}")
    #             return
    #
    #         if provider == "Outlook":
    #             server = smtplib.SMTP(smtp_server, smtp_port)
    #             server.starttls()
    #         else:
    #             server = smtplib.SMTP_SSL(smtp_server, smtp_port)
    #
    #         server.login(email, password)
    #
    #         def read_word_template(file_path):
    #             try:
    #                 doc = Document(file_path)
    #                 html = ""
    #                 for paragraph in doc.paragraphs:
    #                     html += '<p>'
    #                     for run in paragraph.runs:
    #                         if run.bold:
    #                             html += '<b>'
    #                         if run.italic:
    #                             html += '<i>'
    #                         if run.underline:
    #                             html += '<u>'
    #                         if run.font.color.rgb:
    #                             color = run.font.color.rgb
    #                             html += f'<span style="color:#{color};'
    #
    #                             # Check if run has highlight color
    #                             if hasattr(run.font, 'highlight') and run.font.highlight.color:
    #                                 highlight_color = run.font.highlight.color.rgb
    #                                 html += f' background-color:#{highlight_color};'
    #                             else:
    #                                 html += ' background-color:#FF0000;'  # Default to red if highlight color not supported
    #
    #                             html += '">'
    #                         html += run.text
    #                         if run.font.color.rgb or (hasattr(run.font, 'highlight') and run.font.highlight.color):
    #                             html += '</span>'
    #                         if run.underline:
    #                             html += '</u>'
    #                         if run.italic:
    #                             html += '</i>'
    #                         if run.bold:
    #                             html += '</b>'
    #                     html += '</p>'
    #                 return html
    #             except Exception as e:
    #                 print(f"Error reading Word template: {e}")
    #                 # Handle the error or notify the user appropriately
    #                 return ""
    #
    #         def replace_content(html_content, customized_message):
    #             try:
    #                 placeholder = '[[CUSTOM_CONTENT]]'
    #                 html_content = html_content.replace(placeholder, customized_message)
    #                 return html_content
    #             except Exception as e:
    #                 print(f"Error replacing content in HTML template: {e}")
    #                 return html_content
    #
    #         def replace_variables(template, row):
    #             try:
    #                 pattern = re.compile(r'<<([^<>]+)>>')
    #                 matches = pattern.findall(template)
    #                 for match in matches:
    #                     if match in row:
    #                         value = row[match]
    #                         if pd.isna(value):
    #                             value = ""
    #                         template = template.replace(f'<<{match}>>', str(value))
    #                     else:
    #                         print(f"Warning: Variable '{match}' not found in the Excel file")
    #                 return template
    #             except Exception as e:
    #                 print(f"Error replacing variables: {e}")
    #                 return template
    #
    #         def read_html_template(file_path):
    #             try:
    #                 with open(file_path, 'r', encoding='utf-8') as file:
    #                     html_content = file.read()
    #                     return html_content
    #             except Exception as e:
    #                 print(f"Error reading HTML template: {e}")
    #                 return ""
    #
    #         df = pd.read_excel(excel_path)
    #
    #         for index, row in df.iterrows():
    #             try:
    #                 recipient_email = row.get('email')
    #                 if not recipient_email:
    #                     print(f"Skipping row {index + 1}: No email specified")
    #                     continue
    #
    #                 attachments = row.get('attachments')
    #                 if pd.isna(attachments):
    #                     attachments = ''
    #
    #                 # Customize subject using variables from Excel
    #                 subject_template = self.entry_subject.get()
    #                 customized_subject = replace_variables(subject_template, row)
    #
    #                 # Read and customize the Word template
    #                 message_template = read_word_template(word_path)
    #                 customized_message = replace_variables(message_template, row)
    #
    #                 html_content = read_html_template(html_path)
    #                 html_content = replace_content(html_content, customized_message)
    #
    #                 msg = MIMEMultipart()
    #                 msg['From'] = email
    #                 msg['To'] = recipient_email
    #                 msg['Subject'] = customized_subject
    #                 msg.attach(MIMEText(html_content, 'html'))
    #
    #                 if attachments:
    #                     attachment_files = [attachment.strip() for attachment in attachments.split(',')]
    #                     for attachment_file in attachment_files:
    #                         attachment_path = os.path.join('attachments', attachment_file)
    #                         if os.path.exists(attachment_path):
    #                             with open(attachment_path, "rb") as attachment:
    #                                 part = MIMEBase('application', 'octet-stream')
    #                                 part.set_payload(attachment.read())
    #                                 encoders.encode_base64(part)
    #                                 part.add_header('Content-Disposition',
    #                                                 f'attachment; filename={os.path.basename(attachment_path)}')
    #                                 msg.attach(part)
    #                         else:
    #                             print(f"Warning: Attachment '{attachment_file}' not found")
    #
    #                 text = msg.as_string()
    #                 server.sendmail(email, recipient_email, text)
    #
    #             except Exception as e:
    #                 print(f"Error sending email to {recipient_email}: {e}")
    #
    #         server.quit()
    #         messagebox.showinfo("Success", "Emails sent successfully")
    #
    #     except Exception as e:
    #         print(f"Error: {e}")
    #         messagebox.showerror("Error", "Failed to send emails")

    def send_emails(self):
        try:
            word_path = self.word_template_path
            html_path = self.html_template_path
            excel_path = self.excel_file_path
            email = self.selected_email.get()
            password, provider = self.email_credentials.get(email, ("", ""))  # Get password and provider tuple

            # Validate input
            if not word_path or not os.path.exists(word_path):
                messagebox.showerror("Error", "Please choose a Word template file")
                return

            if not excel_path or not os.path.exists(excel_path):
                messagebox.showerror("Error", "Please choose an Excel file")
                return

            if not email or not password:
                messagebox.showerror("Error", "Please choose a valid email account")
                return

            # Set SMTP server and port based on provider
            if provider == "Gmail":
                smtp_server = 'smtp.gmail.com'
                smtp_port = 465
            elif provider == "Outlook":
                smtp_server = 'smtp-mail.outlook.com'
                smtp_port = 587
            elif provider == "Webmail":
                smtp_server = 'mail.una-oic.org'
                smtp_port = 465
            else:
                print("Error", f"Unknown email provider: {provider}")
                return

            if provider == "Outlook":
                server = smtplib.SMTP(smtp_server, smtp_port)
                server.starttls()
            else:
                server = smtplib.SMTP_SSL(smtp_server, smtp_port)

            server.login(email, password)

            def read_word_template(file_path):
                try:
                    doc = Document(file_path)
                    html = ""
                    for paragraph in doc.paragraphs:
                        html += '<p>'
                        for run in paragraph.runs:
                            if run.bold:
                                html += '<b>'
                            if run.italic:
                                html += '<i>'
                            if run.underline:
                                html += '<u>'
                            if run.font.color.rgb:
                                color = run.font.color.rgb
                                html += f'<span style="color:#{color};'

                                # Check if run has highlight color
                                if hasattr(run.font, 'highlight') and run.font.highlight.color:
                                    highlight_color = run.font.highlight.color.rgb
                                    html += f' background-color:#{highlight_color};'
                                else:
                                    html += ' background-color:#FF0000;'  # Default to red if highlight color not supported

                                html += '">'
                            html += run.text
                            if run.font.color.rgb or (hasattr(run.font, 'highlight') and run.font.highlight.color):
                                html += '</span>'
                            if run.underline:
                                html += '</u>'
                            if run.italic:
                                html += '</i>'
                            if run.bold:
                                html += '</b>'
                        html += '</p>'
                    return html
                except Exception as e:
                    print(f"Error reading Word template: {e}")
                    # Handle the error or notify the user appropriately
                    return ""

            def replace_content(html_content, customized_message):
                try:
                    placeholder = '[[CUSTOM_CONTENT]]'
                    html_content = html_content.replace(placeholder, customized_message)
                    return html_content
                except Exception as e:
                    print(f"Error replacing content in HTML template: {e}")
                    return html_content

            def replace_variables(template, row):
                try:
                    pattern = re.compile(r'<<([^<>]+)>>')
                    matches = pattern.findall(template)
                    for match in matches:
                        if match in row:
                            value = row[match]
                            if pd.isna(value):
                                value = ""
                            template = template.replace(f'<<{match}>>', str(value))
                        else:
                            print(f"Warning: Variable '{match}' not found in the Excel file")
                    return template
                except Exception as e:
                    print(f"Error replacing variables: {e}")
                    return template

            def read_html_template(file_path, data_row):
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        html_content = file.read()

                        # Replace variables in HTML template with values from data_row
                        pattern = re.compile(r'<<([^<>]+)>>')
                        matches = pattern.findall(html_content)
                        for match in matches:
                            if match in data_row:
                                value = data_row[match]
                                if pd.isna(value):
                                    value = ""
                                html_content = html_content.replace(f'<<{match}>>', str(value))
                            else:
                                print(f"Warning: Variable '{match}' not found in the Excel file")

                        return html_content
                except Exception as e:
                    print(f"Error reading HTML template: {e}")
                    return ""

            df = pd.read_excel(excel_path)

            for index, row in df.iterrows():
                try:
                    recipient_email = row.get('email')
                    if not recipient_email:
                        print(f"Skipping row {index + 1}: No email specified")
                        continue

                    attachments = row.get('attachments')
                    if pd.isna(attachments):
                        attachments = ''

                    # Customize subject using variables from Excel
                    subject_template = self.entry_subject.get()
                    customized_subject = replace_variables(subject_template, row)

                    # Read and customize the Word template
                    message_template = read_word_template(word_path)
                    customized_message = replace_variables(message_template, row)

                    # Read and customize the HTML template
                    html_content = read_html_template(html_path, row)
                    html_content = replace_content(html_content, customized_message)

                    msg = MIMEMultipart()
                    msg['From'] = email
                    msg['To'] = recipient_email
                    msg['Subject'] = customized_subject
                    msg.attach(MIMEText(html_content, 'html'))

                    if attachments:
                        attachment_files = [attachment.strip() for attachment in attachments.split(',')]
                        for attachment_file in attachment_files:
                            attachment_path = os.path.join('attachments', attachment_file)
                            if os.path.exists(attachment_path):
                                with open(attachment_path, "rb") as attachment:
                                    part = MIMEBase('application', 'octet-stream')
                                    part.set_payload(attachment.read())
                                    encoders.encode_base64(part)
                                    part.add_header('Content-Disposition',
                                                    f'attachment; filename={os.path.basename(attachment_path)}')
                                    msg.attach(part)
                            else:
                                print(f"Warning: Attachment '{attachment_file}' not found")

                    text = msg.as_string()
                    server.sendmail(email, recipient_email, text)

                except Exception as e:
                    print(f"Error sending email to {recipient_email}: {e}")

            server.quit()
            messagebox.showinfo("Success", "Emails sent successfully")

        except Exception as e:
            print(f"Error: {e}")
            messagebox.showerror("Error", "Failed to send emails")


if __name__ == "__main__":
    app = EmailSenderApp()
    app.mainloop()
