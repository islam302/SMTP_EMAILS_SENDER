# # import re
# # import pandas as pd
# # from docx import Document
# # from docx2pdf import convert
# # from tkinter import *
# # from tkinter import ttk, messagebox, filedialog
# # from PIL import Image, ImageTk
# # import os
# # import sys
# # import io
# # import glob
# # class Replace(Tk):
# #
# #     def __init__(self):
# #         super().__init__()
# #         self.title("Create Pdf files")
# #         self.geometry("800x600")
# #         self.configure(bg="#F0F0F0")
# #
# #         self.load_background_image()
# #
# #         self.style = ttk.Style()
# #         self.style.theme_use("clam")
# #         self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 12, 'bold'))
# #         self.style.map('TButton', background=[('active', '#005A9E')])
# #         self.style.configure('TLabel', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
# #         self.style.configure('TEntry', font=('Segoe UI', 12))
# #         self.style.configure('TCheckbutton', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
# #         self.style.map('TCheckbutton', background=[('active', '#005A9E')], foreground=[('active', 'white')])
# #
# #         self.style.configure('Red.TButton', foreground='white', background='red', font=('Segoe UI', 12, 'bold'))
# #         self.style.map('Red.TButton', background=[('active', '#8B0000')])
# #
# #         self.word_template_path = None
# #         self.excel_file_path = None
# #
# #         self.create_widgets()
# #
# #     def load_background_image(self):
# #         current_dir = os.path.dirname(sys.argv[0])
# #         background_files = glob.glob(os.path.join(current_dir, "logo.*"))
# #         if background_files:
# #             background_image_path = background_files[0]
# #             if os.path.exists(background_image_path):
# #                 try:
# #                     background_image = Image.open(background_image_path)
# #                     background_image = background_image.resize((800, 600), Image.LANCZOS)
# #                     self.background_photo = ImageTk.PhotoImage(background_image)
# #                     self.canvas_bg = Canvas(self, width=800, height=600)
# #                     self.canvas_bg.create_image(0, 0, anchor='nw', image=self.background_photo)
# #                     self.canvas_bg.pack(fill='both', expand=True)
# #                 except Exception:
# #                     messagebox.showerror("Error", "Failed to load background image")
# #         else:
# #             print("Background image not found")
# #
# #     def create_widgets(self):
# #         btn_browse_word = ttk.Button(self, text="WORD FILE", command=self.browse_word_template)
# #         self.canvas_bg.create_window(400, 160, anchor='center', window=btn_browse_word)
# #
# #         btn_browse_excel = ttk.Button(self, text="EXCEL FILE", command=self.browse_excel_file)
# #         self.canvas_bg.create_window(400, 220, anchor='center', window=btn_browse_excel)
# #
# #         self.word_var = BooleanVar()
# #         self.pdf_var = BooleanVar()
# #
# #         chk_word = ttk.Checkbutton(self, text="Create Word File", variable=self.word_var)
# #         self.canvas_bg.create_window(340, 280, anchor='center', window=chk_word)
# #
# #         chk_pdf = ttk.Checkbutton(self, text="Create PDF File", variable=self.pdf_var)
# #         self.canvas_bg.create_window(480, 280, anchor='center', window=chk_pdf)
# #
# #         btn_generate_files = ttk.Button(self, text="Create Files", command=self.generate_files)
# #         self.canvas_bg.create_window(400, 340, anchor='center', window=btn_generate_files)
# #
# #         btn_exit = ttk.Button(self, text="EXIT", style='Red.TButton', command=self.destroy)
# #         self.canvas_bg.create_window(400, 400, anchor='center', window=btn_exit)
# #
# #     def browse_word_template(self):
# #         self.word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
# #         if self.word_template_path:
# #             messagebox.showinfo("Selected File", f"Word Template: {self.word_template_path}")
# #
# #     def browse_excel_file(self):
# #         self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
# #         if self.excel_file_path:
# #             messagebox.showinfo("Selected File", f"Excel File: {self.excel_file_path}")
# #
# #     def generate_files(self):
# #         if not self.word_template_path or not self.excel_file_path:
# #             messagebox.showerror("Error", "Please select both a Word template and an Excel file.")
# #             return
# #
# #         if not (self.word_var.get() or self.pdf_var.get()):
# #             messagebox.showerror("Error", "Please select at least one file type to create.")
# #             return
# #
# #         # Create results directory if it doesn't exist
# #         results_dir = "results"
# #         if not os.path.exists(results_dir):
# #             os.makedirs(results_dir)
# #
# #         create_word_and_pdf(self.word_template_path, self.excel_file_path, results_dir, self.word_var.get(),
# #                             self.pdf_var.get())
# #         messagebox.showinfo("Success", "Files generated successfully")
# # def replace_variables(paragraph, row):
# #     pattern = re.compile(r'<<([^<>]+)>>')
# #     full_text = ''.join(run.text for run in paragraph.runs)  # Combine all text from runs
# #     updated_text = full_text
# #
# #     # Replace variables in the combined text
# #     for match in pattern.findall(full_text):
# #         key = match.strip()
# #         if key in row:
# #             value = row[key]
# #             if pd.isna(value):
# #                 value = ""
# #             updated_text = updated_text.replace(f'<<{key}>>', str(value))
# #         else:
# #             print(f"Warning: Variable '{key}' not found in the Excel file")
# #
# #     # Update runs with the new text
# #     current_pos = 0
# #     for run in paragraph.runs:
# #         run_length = len(run.text)
# #         run.text = updated_text[current_pos:current_pos + run_length]
# #         current_pos += run_length
# #
# #         # If we run out of characters in updated_text, empty the remaining runs
# #         if current_pos >= len(updated_text):
# #             break
# #
# #     # If there are remaining characters in updated_text, add them to the last run
# #     if current_pos < len(updated_text):
# #         paragraph.runs[-1].text += updated_text[current_pos:]
# # def create_word_and_pdf(template_path, excel_path, output_dir, create_word, create_pdf):
# #     doc = Document(template_path)
# #     df = pd.read_excel(excel_path)
# #
# #     for idx, row in df.iterrows():
# #         new_doc = Document(template_path)
# #
# #         for paragraph in new_doc.paragraphs:
# #             replace_variables(paragraph, row)
# #
# #         for section in new_doc.sections:
# #             for header in section.header.paragraphs:
# #                 replace_variables(header, row)
# #             for footer in section.footer.paragraphs:
# #                 replace_variables(footer, row)
# #
# #         base_filename = str(row.get('file name', f"output_document_{idx + 1}"))
# #
# #         if create_word:
# #             new_word_path = os.path.join(output_dir, f"{base_filename}.docx")
# #             new_doc.save(new_word_path)
# #
# #         if create_pdf:
# #             if create_word:
# #                 sys.stdout = io.StringIO()
# #                 convert(new_word_path, os.path.join(output_dir, f"{base_filename}.pdf"))
# #                 sys.stdout = sys.__stdout__
# #             else:
# #                 new_temp_word_path = os.path.join(output_dir, f"temp_{base_filename}.docx")
# #                 new_doc.save(new_temp_word_path)
# #                 sys.stdout = io.StringIO()
# #                 convert(new_temp_word_path, os.path.join(output_dir, f"{base_filename}.pdf"))
# #                 sys.stdout = sys.__stdout__
# #                 os.remove(new_temp_word_path)
# # if __name__ == "__main__":
# #     app = Replace()
# #     app.protocol("WM_DELETE_WINDOW", app.destroy)
# #     app.mainloop()
# #
#
# #
# # import re
# # import pandas as pd
# # from docx import Document
# # from docx2pdf import convert
# # from tkinter import *
# # from tkinter import ttk, messagebox, filedialog
# # from PIL import Image, ImageTk
# # import os
# # import sys
# # import io
# # import glob
# #
# # class Replace(Tk):
# #     def __init__(self):
# #         super().__init__()
# #         self.title("Create Pdf files")
# #         self.geometry("800x600")
# #         self.configure(bg="#F0F0F0")
# #
# #         self.load_background_image()
# #
# #         self.style = ttk.Style()
# #         self.style.theme_use("clam")
# #         self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 12, 'bold'))
# #         self.style.map('TButton', background=[('active', '#005A9E')])
# #         self.style.configure('TLabel', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
# #         self.style.configure('TEntry', font=('Segoe UI', 12))
# #         self.style.configure('TCheckbutton', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
# #         self.style.map('TCheckbutton', background=[('active', '#005A9E')], foreground=[('active', 'white')])
# #
# #         self.style.configure('Red.TButton', foreground='white', background='red', font=('Segoe UI', 12, 'bold'))
# #         self.style.map('Red.TButton', background=[('active', '#8B0000')])
# #
# #         self.word_template_path = None
# #         self.excel_file_path = None
# #
# #         self.create_widgets()
# #
# #     def load_background_image(self):
# #         current_dir = os.path.dirname(sys.argv[0])
# #         background_files = glob.glob(os.path.join(current_dir, "logo.*"))
# #         if background_files:
# #             background_image_path = background_files[0]
# #             if os.path.exists(background_image_path):
# #                 try:
# #                     background_image = Image.open(background_image_path)
# #                     background_image = background_image.resize((800, 600), Image.LANCZOS)
# #                     self.background_photo = ImageTk.PhotoImage(background_image)
# #                     self.canvas_bg = Canvas(self, width=800, height=600)
# #                     self.canvas_bg.create_image(0, 0, anchor='nw', image=self.background_photo)
# #                     self.canvas_bg.pack(fill='both', expand=True)
# #                 except Exception:
# #                     messagebox.showerror("Error", "Failed to load background image")
# #         else:
# #             print("Background image not found")
# #
# #     def create_widgets(self):
# #         btn_browse_word = ttk.Button(self, text="WORD FILE", command=self.browse_word_template)
# #         self.canvas_bg.create_window(400, 160, anchor='center', window=btn_browse_word)
# #
# #         btn_browse_excel = ttk.Button(self, text="EXCEL FILE", command=self.browse_excel_file)
# #         self.canvas_bg.create_window(400, 220, anchor='center', window=btn_browse_excel)
# #
# #         self.word_var = BooleanVar()
# #         self.pdf_var = BooleanVar()
# #
# #         chk_word = ttk.Checkbutton(self, text="Create Word File", variable=self.word_var)
# #         self.canvas_bg.create_window(340, 280, anchor='center', window=chk_word)
# #
# #         chk_pdf = ttk.Checkbutton(self, text="Create PDF File", variable=self.pdf_var)
# #         self.canvas_bg.create_window(480, 280, anchor='center', window=chk_pdf)
# #
# #         btn_generate_files = ttk.Button(self, text="Create Files", command=self.generate_files)
# #         self.canvas_bg.create_window(400, 340, anchor='center', window=btn_generate_files)
# #
# #         btn_exit = ttk.Button(self, text="EXIT", style='Red.TButton', command=self.destroy)
# #         self.canvas_bg.create_window(400, 400, anchor='center', window=btn_exit)
# #
# #     def browse_word_template(self):
# #         self.word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
# #         if self.word_template_path:
# #             messagebox.showinfo("Selected File", f"Word Template: {self.word_template_path}")
# #
# #     def browse_excel_file(self):
# #         self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
# #         if self.excel_file_path:
# #             messagebox.showinfo("Selected File", f"Excel File: {self.excel_file_path}")
# #
# #     def generate_files(self):
# #         if not self.word_template_path or not self.excel_file_path:
# #             messagebox.showerror("Error", "Please select both a Word template and an Excel file.")
# #             return
# #
# #         if not (self.word_var.get() or self.pdf_var.get()):
# #             messagebox.showerror("Error", "Please select at least one file type to create.")
# #             return
# #
# #         # Create results directory if it doesn't exist
# #         results_dir = "results"
# #         if not os.path.exists(results_dir):
# #             os.makedirs(results_dir)
# #
# #         create_word_and_pdf(self.word_template_path, self.excel_file_path, results_dir, self.word_var.get(), self.pdf_var.get())
# #         messagebox.showinfo("Success", "Files generated successfully")
# #
# # def replace_variables(paragraph, row):
# #     pattern = re.compile(r'<<([^<>]+)>>')
# #     for match in pattern.findall(paragraph.text):
# #         key = match.strip()
# #         if key in row:
# #             value = row[key]
# #             if pd.isna(value):
# #                 value = ""
# #             paragraph.text = paragraph.text.replace(f'<<{key}>>', str(value))
# #         else:
# #             print(f"Warning: Variable '{key}' not found in the Excel file")
# #
# # def copy_media(src_doc, dest_doc):
# #     for rel in src_doc.part.rels.values():
# #         if "image" in rel.target_ref:
# #             dest_doc.part.rels.add_relationship(
# #                 rel.reltype,
# #                 rel.target_ref,
# #                 rel.rId,
# #                 is_external=rel.is_external
# #             )
# #
# # def create_word_and_pdf(template_path, excel_path, output_dir, create_word, create_pdf):
# #     doc = Document(template_path)
# #     df = pd.read_excel(excel_path)
# #
# #     for idx, row in df.iterrows():
# #         new_doc = Document(template_path)
# #
# #         for paragraph in new_doc.paragraphs:
# #             replace_variables(paragraph, row)
# #
# #         for section in new_doc.sections:
# #             for header in section.header.paragraphs:
# #                 replace_variables(header, row)
# #             for footer in section.footer.paragraphs:
# #                 replace_variables(footer, row)
# #
# #         copy_media(doc, new_doc)
# #
# #         base_filename = str(row.get('file name', f"output_document_{idx + 1}"))
# #
# #         if create_word:
# #             new_word_path = os.path.join(output_dir, f"{base_filename}.docx")
# #             new_doc.save(new_word_path)
# #
# #         if create_pdf:
# #             if create_word:
# #                 sys.stdout = io.StringIO()
# #                 convert(new_word_path, os.path.join(output_dir, f"{base_filename}.pdf"))
# #                 sys.stdout = sys.__stdout__
# #             else:
# #                 new_temp_word_path = os.path.join(output_dir, f"temp_{base_filename}.docx")
# #                 new_doc.save(new_temp_word_path)
# #                 sys.stdout = io.StringIO()
# #                 convert(new_temp_word_path, os.path.join(output_dir, f"{base_filename}.pdf"))
# #                 sys.stdout = sys.__stdout__
# #                 os.remove(new_temp_word_path)
# #
# # if __name__ == "__main__":
# #     app = Replace()
# #     app.protocol("WM_DELETE_WINDOW", app.destroy)
# #     app.mainloop()
#
#
#
# import os
# import re
# import pandas as pd
# from docx import Document
# from docx2pdf import convert
# from tkinter import *
# from tkinter import ttk, messagebox, filedialog
# from PIL import Image, ImageTk
# import io
# import glob
# from docx.shared import Inches
# import sys
#
# class Replace(Tk):
#
#     def __init__(self):
#         super().__init__()
#         self.title("Create Pdf files")
#         self.geometry("800x600")
#         self.configure(bg="#F0F0F0")
#
#         self.load_background_image()
#
#         self.style = ttk.Style()
#         self.style.theme_use("clam")
#         self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 12, 'bold'))
#         self.style.map('TButton', background=[('active', '#005A9E')])
#         self.style.configure('TLabel', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
#         self.style.configure('TEntry', font=('Segoe UI', 12))
#         self.style.configure('TCheckbutton', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
#         self.style.map('TCheckbutton', background=[('active', '#005A9E')], foreground=[('active', 'white')])
#
#         self.style.configure('Red.TButton', foreground='white', background='red', font=('Segoe UI', 12, 'bold'))
#         self.style.map('Red.TButton', background=[('active', '#8B0000')])
#
#         self.word_template_path = None
#         self.excel_file_path = None
#
#         self.create_widgets()
#
#     def load_background_image(self):
#         current_dir = os.path.dirname(sys.argv[0])
#         background_files = glob.glob(os.path.join(current_dir, "logo.*"))
#         if background_files:
#             background_image_path = background_files[0]
#             if os.path.exists(background_image_path):
#                 try:
#                     background_image = Image.open(background_image_path)
#                     background_image = background_image.resize((800, 600), Image.LANCZOS)
#                     self.background_photo = ImageTk.PhotoImage(background_image)
#                     self.canvas_bg = Canvas(self, width=800, height=600)
#                     self.canvas_bg.create_image(0, 0, anchor='nw', image=self.background_photo)
#                     self.canvas_bg.pack(fill='both', expand=True)
#                 except Exception:
#                     messagebox.showerror("Error", "Failed to load background image")
#         else:
#             print("Background image not found")
#
#     def create_widgets(self):
#         btn_browse_word = ttk.Button(self, text="WORD FILE", command=self.browse_word_template)
#         self.canvas_bg.create_window(400, 160, anchor='center', window=btn_browse_word)
#
#         btn_browse_excel = ttk.Button(self, text="EXCEL FILE", command=self.browse_excel_file)
#         self.canvas_bg.create_window(400, 220, anchor='center', window=btn_browse_excel)
#
#         self.word_var = BooleanVar()
#         self.pdf_var = BooleanVar()
#
#         chk_word = ttk.Checkbutton(self, text="Create Word File", variable=self.word_var)
#         self.canvas_bg.create_window(340, 280, anchor='center', window=chk_word)
#
#         chk_pdf = ttk.Checkbutton(self, text="Create PDF File", variable=self.pdf_var)
#         self.canvas_bg.create_window(480, 280, anchor='center', window=chk_pdf)
#
#         btn_generate_files = ttk.Button(self, text="Create Files", command=self.generate_files)
#         self.canvas_bg.create_window(400, 340, anchor='center', window=btn_generate_files)
#
#         btn_exit = ttk.Button(self, text="EXIT", style='Red.TButton', command=self.destroy)
#         self.canvas_bg.create_window(400, 400, anchor='center', window=btn_exit)
#
#     def browse_word_template(self):
#         self.word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
#         if self.word_template_path:
#             messagebox.showinfo("Selected File", f"Word Template: {self.word_template_path}")
#
#     def browse_excel_file(self):
#         self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
#         if self.excel_file_path:
#             messagebox.showinfo("Selected File", f"Excel File: {self.excel_file_path}")
#
#     def generate_files(self):
#         if not self.word_template_path or not self.excel_file_path:
#             messagebox.showerror("Error", "Please select both a Word template and an Excel file.")
#             return
#
#         if not (self.word_var.get() or self.pdf_var.get()):
#             messagebox.showerror("Error", "Please select at least one file type to create.")
#             return
#
#         # Create results directory if it doesn't exist
#         results_dir = "results"
#         if not os.path.exists(results_dir):
#             os.makedirs(results_dir)
#
#         create_word_and_pdf(self.word_template_path, self.excel_file_path, results_dir, self.word_var.get(),
#                             self.pdf_var.get())
#         messagebox.showinfo("Success", "Files generated successfully")
#
#
# images_dir = "images"
#
# def replace_variables(paragraph, row):
#     pattern = re.compile(r'<<([^<>]+)>>')
#     full_text = ''.join(run.text for run in paragraph.runs)  # Combine all text from runs
#     updated_text = full_text
#
#     # Replace variables in the combined text
#     for match in pattern.findall(full_text):
#         key = match.strip()
#         if key in row:
#             value = row[key]
#             if pd.isna(value):
#                 value = ""
#             updated_text = updated_text.replace(f'<<{key}>>', str(value))
#         else:
#             print(f"Warning: Variable '{key}' not found in the Excel file")
#
#     # Update runs with the new text
#     current_pos = 0
#     for run in paragraph.runs:
#         run_length = len(run.text)
#         run.text = updated_text[current_pos:current_pos + run_length]
#         current_pos += run_length
#
#         # If we run out of characters in updated_text, empty the remaining runs
#         if current_pos >= len(updated_text):
#             break
#
#     # If there are remaining characters in updated_text, add them to the last run
#     if current_pos < len(updated_text):
#         paragraph.runs[-1].text += updated_text[current_pos:]
#
#
#
# def insert_image(doc, images_dir):
#     """
#     Inserts images into the Word document where placeholders are found.
#
#     Args:
#         doc (Document): The Document object where images will be inserted.
#         images_dir (str): The directory containing the images.
#     """
#     for paragraph in doc.paragraphs:
#         # Find all placeholders for images in the paragraph text
#         image_placeholders = re.findall(r'\[image([^\]]+)\]', paragraph.text)
#
#         for image_name in image_placeholders:
#             image_name = image_name.strip()
#             image_file = os.path.abspath(os.path.join(images_dir, image_name))
#
#             if os.path.exists(image_file):
#                 # Create a new run for the image
#                 run = paragraph.add_run()
#                 run.add_picture(image_file, width=Inches(2), height=Inches(2))
#
#                 # Replace the placeholder with an empty string
#                 paragraph.text = paragraph.text.replace(f'[image{image_name}]', '')
#             else:
#                 print(f"Warning: Image file not found: {image_file}")
#
#     # Remove any remaining placeholders
#     for paragraph in doc.paragraphs:
#         paragraph.text = re.sub(r'\[image[^\]]+\]', '', paragraph.text)
#
#
# def create_word_and_pdf(word_template_path, excel_file_path, results_dir, create_word, create_pdf):
#     df = pd.read_excel(excel_file_path)
#
#     for index, row in df.iterrows():
#         doc = Document(word_template_path)
#
#         for paragraph in doc.paragraphs:
#             replace_variables(paragraph, row)
#
#         insert_image(doc, "images")
#
#         file_name = f"{row['email']}.docx"
#         doc_path = os.path.join(results_dir, file_name)
#         doc.save(doc_path)
#
#         if create_pdf:
#             pdf_path = os.path.splitext(doc_path)[0] + '.pdf'
#             convert(doc_path, pdf_path)
#
# if __name__ == "__main__":
#     app = Replace()
#     app.mainloop()


# import os
# import shutil
# import pandas as pd
# from docx import Document
# from docx2pdf import convert
# from tkinter import *
# from tkinter import ttk, messagebox, filedialog
# from PIL import Image, ImageTk
# import glob
# import sys
#
# class Replace(Tk):
#
#     def __init__(self):
#         super().__init__()
#         self.title("Create Pdf files")
#         self.geometry("800x600")
#         self.configure(bg="#F0F0F0")
#
#         self.load_background_image()
#
#         self.style = ttk.Style()
#         self.style.theme_use("clam")
#         self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 12, 'bold'))
#         self.style.map('TButton', background=[('active', '#005A9E')])
#         self.style.configure('TLabel', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
#         self.style.configure('TEntry', font=('Segoe UI', 12))
#         self.style.configure('TCheckbutton', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
#         self.style.map('TCheckbutton', background=[('active', '#005A9E')], foreground=[('active', 'white')])
#
#         self.style.configure('Red.TButton', foreground='white', background='red', font=('Segoe UI', 12, 'bold'))
#         self.style.map('Red.TButton', background=[('active', '#8B0000')])
#
#         self.word_template_path = None
#         self.excel_file_path = None
#
#         self.create_widgets()
#
#     def load_background_image(self):
#         current_dir = os.path.dirname(sys.argv[0])
#         background_files = glob.glob(os.path.join(current_dir, "logo.*"))
#         if background_files:
#             background_image_path = background_files[0]
#             if os.path.exists(background_image_path):
#                 try:
#                     background_image = Image.open(background_image_path)
#                     background_image = background_image.resize((800, 600), Image.LANCZOS)
#                     self.background_photo = ImageTk.PhotoImage(background_image)
#                     self.canvas_bg = Canvas(self, width=800, height=600)
#                     self.canvas_bg.create_image(0, 0, anchor='nw', image=self.background_photo)
#                     self.canvas_bg.pack(fill='both', expand=True)
#                 except Exception:
#                     messagebox.showerror("Error", "Failed to load background image")
#         else:
#             print("Background image not found")
#
#     def create_widgets(self):
#         btn_browse_word = ttk.Button(self, text="WORD FILE", command=self.browse_word_template)
#         self.canvas_bg.create_window(400, 160, anchor='center', window=btn_browse_word)
#
#         btn_browse_excel = ttk.Button(self, text="EXCEL FILE", command=self.browse_excel_file)
#         self.canvas_bg.create_window(400, 220, anchor='center', window=btn_browse_excel)
#
#         self.word_var = BooleanVar()
#         self.pdf_var = BooleanVar()
#
#         chk_word = ttk.Checkbutton(self, text="Create Word File", variable=self.word_var)
#         self.canvas_bg.create_window(340, 280, anchor='center', window=chk_word)
#
#         chk_pdf = ttk.Checkbutton(self, text="Create PDF File", variable=self.pdf_var)
#         self.canvas_bg.create_window(480, 280, anchor='center', window=chk_pdf)
#
#         btn_generate_files = ttk.Button(self, text="Create Files", command=self.generate_files)
#         self.canvas_bg.create_window(400, 340, anchor='center', window=btn_generate_files)
#
#         btn_exit = ttk.Button(self, text="EXIT", style='Red.TButton', command=self.destroy)
#         self.canvas_bg.create_window(400, 400, anchor='center', window=btn_exit)
#
#     def browse_word_template(self):
#         self.word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
#         if self.word_template_path:
#             messagebox.showinfo("Selected File", f"Word Template: {self.word_template_path}")
#
#     def browse_excel_file(self):
#         self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
#         if self.excel_file_path:
#             messagebox.showinfo("Selected File", f"Excel File: {self.excel_file_path}")
#
#     def generate_files(self):
#         if not self.word_template_path or not self.excel_file_path:
#             messagebox.showerror("Error", "Please select both a Word template and an Excel file.")
#             return
#
#         if not (self.word_var.get() or self.pdf_var.get()):
#             messagebox.showerror("Error", "Please select at least one file type to create.")
#             return
#
#         # Create results directory if it doesn't exist
#         results_dir = "results"
#         if not os.path.exists(results_dir):
#             os.makedirs(results_dir)
#
#         # Copy the Word template to the results directory
#         shutil.copy(self.word_template_path, os.path.join(results_dir, "template.docx"))
#
#         # If the PDF option is selected, also convert the copied Word file to PDF
#         if self.pdf_var.get():
#             doc_path = os.path.join(results_dir, "template.docx")
#             pdf_path = os.path.splitext(doc_path)[0] + '.pdf'
#             convert(doc_path, pdf_path)
#
#         messagebox.showinfo("Success", "Files generated successfully")
#
# if __name__ == "__main__":
#     app = Replace()
#     app.mainloop()


import os
import shutil
import pandas as pd
from docx import Document
from docx2pdf import convert
from tkinter import *
from tkinter import ttk, messagebox, filedialog
from PIL import Image, ImageTk
import glob
import sys

class Replace(Tk):

    def __init__(self):
        super().__init__()
        self.title("Create Pdf files")
        self.geometry("800x600")
        self.configure(bg="#F0F0F0")

        self.load_background_image()

        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 12, 'bold'))
        self.style.map('TButton', background=[('active', '#005A9E')])
        self.style.configure('TLabel', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
        self.style.configure('TEntry', font=('Segoe UI', 12))
        self.style.configure('TCheckbutton', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
        self.style.map('TCheckbutton', background=[('active', '#005A9E')], foreground=[('active', 'white')])

        self.style.configure('Red.TButton', foreground='white', background='red', font=('Segoe UI', 12, 'bold'))
        self.style.map('Red.TButton', background=[('active', '#8B0000')])

        self.word_template_path = None
        self.excel_file_path = None

        self.create_widgets()

    def load_background_image(self):
        current_dir = os.path.dirname(sys.argv[0])
        background_files = glob.glob(os.path.join(current_dir, "logo.*"))
        if background_files:
            background_image_path = background_files[0]
            if os.path.exists(background_image_path):
                try:
                    background_image = Image.open(background_image_path)
                    background_image = background_image.resize((800, 600), Image.LANCZOS)
                    self.background_photo = ImageTk.PhotoImage(background_image)
                    self.canvas_bg = Canvas(self, width=800, height=600)
                    self.canvas_bg.create_image(0, 0, anchor='nw', image=self.background_photo)
                    self.canvas_bg.pack(fill='both', expand=True)
                except Exception:
                    messagebox.showerror("Error", "Failed to load background image")
        else:
            print("Background image not found")

    def create_widgets(self):
        btn_browse_word = ttk.Button(self, text="WORD FILE", command=self.browse_word_template)
        self.canvas_bg.create_window(400, 160, anchor='center', window=btn_browse_word)

        btn_browse_excel = ttk.Button(self, text="EXCEL FILE", command=self.browse_excel_file)
        self.canvas_bg.create_window(400, 220, anchor='center', window=btn_browse_excel)

        self.word_var = BooleanVar()
        self.pdf_var = BooleanVar()

        chk_word = ttk.Checkbutton(self, text="Create Word File", variable=self.word_var)
        self.canvas_bg.create_window(340, 280, anchor='center', window=chk_word)

        chk_pdf = ttk.Checkbutton(self, text="Create PDF File", variable=self.pdf_var)
        self.canvas_bg.create_window(480, 280, anchor='center', window=chk_pdf)

        btn_generate_files = ttk.Button(self, text="Create Files", command=self.generate_files)
        self.canvas_bg.create_window(400, 340, anchor='center', window=btn_generate_files)

        btn_exit = ttk.Button(self, text="EXIT", style='Red.TButton', command=self.destroy)
        self.canvas_bg.create_window(400, 400, anchor='center', window=btn_exit)

    def browse_word_template(self):
        self.word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.word_template_path:
            messagebox.showinfo("Selected File", f"Word Template: {self.word_template_path}")

    def browse_excel_file(self):
        self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if self.excel_file_path:
            messagebox.showinfo("Selected File", f"Excel File: {self.excel_file_path}")

    def generate_files(self):
        if not self.word_template_path or not self.excel_file_path:
            messagebox.showerror("Error", "Please select both a Word template and an Excel file.")
            return

        if not (self.word_var.get() or self.pdf_var.get()):
            messagebox.showerror("Error", "Please select at least one file type to create.")
            return

        # Create results directory if it doesn't exist
        results_dir = "results"
        if not os.path.exists(results_dir):
            os.makedirs(results_dir)

        # Load the Excel file
        df = pd.read_excel(self.excel_file_path)

        # Iterate over rows in the Excel file
        for _, row in df.iterrows():
            file_name = row.get('file name')
            if not file_name:
                continue

            # Copy the Word template to the results directory with the specified filename
            file_path = os.path.join(results_dir, f"{file_name}.docx")
            shutil.copy(self.word_template_path, file_path)

            # Open the copied Word file and replace variables
            doc = Document(file_path)
            for para in doc.paragraphs:
                # Process text runs in paragraphs to replace placeholders without changing formatting
                for run in para.runs:
                    for key, value in row.items():
                        if pd.notna(value):
                            placeholder = f'<<{key}>>'
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))

            # Save the modified document
            doc.save(file_path)

            # Convert to PDF if the option is selected
            if self.pdf_var.get():
                pdf_path = os.path.splitext(file_path)[0] + '.pdf'
                convert(file_path, pdf_path)

        messagebox.showinfo("Success", "Files generated successfully")

if __name__ == "__main__":
    app = Replace()
    app.mainloop()



# import os
# import shutil
# import pandas as pd
# from docx import Document
# from docx2pdf import convert
# from tkinter import *
# from tkinter import ttk, messagebox, filedialog
# from PIL import Image, ImageTk
# import glob
# import sys
#
# class Replace(Tk):
#
#     def __init__(self):
#         super().__init__()
#         self.title("Create Pdf files")
#         self.geometry("800x600")
#         self.configure(bg="#F0F0F0")
#
#         self.load_background_image()
#
#         self.style = ttk.Style()
#         self.style.theme_use("clam")
#         self.style.configure('TButton', foreground='white', background='#0078D7', font=('Segoe UI', 12, 'bold'))
#         self.style.map('TButton', background=[('active', '#005A9E')])
#         self.style.configure('TLabel', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
#         self.style.configure('TEntry', font=('Segoe UI', 12))
#         self.style.configure('TCheckbutton', background='#F0F0F0', foreground='#0078D7', font=('Segoe UI', 12))
#         self.style.map('TCheckbutton', background=[('active', '#005A9E')], foreground=[('active', 'white')])
#
#         self.style.configure('Red.TButton', foreground='white', background='red', font=('Segoe UI', 12, 'bold'))
#         self.style.map('Red.TButton', background=[('active', '#8B0000')])
#
#         self.word_template_path = None
#         self.excel_file_path = None
#
#         self.create_widgets()
#
#     def load_background_image(self):
#         current_dir = os.path.dirname(sys.argv[0])
#         background_files = glob.glob(os.path.join(current_dir, "logo.*"))
#         if background_files:
#             background_image_path = background_files[0]
#             if os.path.exists(background_image_path):
#                 try:
#                     background_image = Image.open(background_image_path)
#                     background_image = background_image.resize((800, 600), Image.LANCZOS)
#                     self.background_photo = ImageTk.PhotoImage(background_image)
#                     self.canvas_bg = Canvas(self, width=800, height=600)
#                     self.canvas_bg.create_image(0, 0, anchor='nw', image=self.background_photo)
#                     self.canvas_bg.pack(fill='both', expand=True)
#                 except Exception:
#                     messagebox.showerror("Error", "Failed to load background image")
#         else:
#             print("Background image not found")
#
#     def create_widgets(self):
#         btn_browse_word = ttk.Button(self, text="WORD FILE", command=self.browse_word_template)
#         self.canvas_bg.create_window(400, 160, anchor='center', window=btn_browse_word)
#
#         btn_browse_excel = ttk.Button(self, text="EXCEL FILE", command=self.browse_excel_file)
#         self.canvas_bg.create_window(400, 220, anchor='center', window=btn_browse_excel)
#
#         self.word_var = BooleanVar()
#         self.pdf_var = BooleanVar()
#
#         chk_word = ttk.Checkbutton(self, text="Create Word File", variable=self.word_var)
#         self.canvas_bg.create_window(340, 280, anchor='center', window=chk_word)
#
#         chk_pdf = ttk.Checkbutton(self, text="Create PDF File", variable=self.pdf_var)
#         self.canvas_bg.create_window(480, 280, anchor='center', window=chk_pdf)
#
#         btn_generate_files = ttk.Button(self, text="Create Files", command=self.generate_files)
#         self.canvas_bg.create_window(400, 340, anchor='center', window=btn_generate_files)
#
#         btn_exit = ttk.Button(self, text="EXIT", style='Red.TButton', command=self.destroy)
#         self.canvas_bg.create_window(400, 400, anchor='center', window=btn_exit)
#
#     def browse_word_template(self):
#         self.word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
#         if self.word_template_path:
#             messagebox.showinfo("Selected File", f"Word Template: {self.word_template_path}")
#
#     def browse_excel_file(self):
#         self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
#         if self.excel_file_path:
#             messagebox.showinfo("Selected File", f"Excel File: {self.excel_file_path}")
#
#     def generate_files(self):
#         if not self.word_template_path or not self.excel_file_path:
#             messagebox.showerror("Error", "Please select both a Word template and an Excel file.")
#             return
#
#         if not (self.word_var.get() or self.pdf_var.get()):
#             messagebox.showerror("Error", "Please select at least one file type to create.")
#             return
#
#         # Create results directory if it doesn't exist
#         results_dir = "results"
#         if not os.path.exists(results_dir):
#             os.makedirs(results_dir)
#
#         # Load the Excel file
#         df = pd.read_excel(self.excel_file_path)
#
#         # Iterate over rows in the Excel file
#         for _, row in df.iterrows():
#             file_name = row.get('file name')
#             if not file_name:
#                 continue
#
#             # Copy the Word template to the results directory with the specified filename
#             file_path = os.path.join(results_dir, f"{file_name}.docx")
#             shutil.copy(self.word_template_path, file_path)
#
#             # Open the copied Word file and replace variables
#             doc = Document(file_path)
#             for para in doc.paragraphs:
#                 # Replace text placeholders in paragraphs
#                 for key, value in row.items():
#                     if pd.notna(value):
#                         placeholder = f'<<{key}>>'
#                         if placeholder in para.text:
#                             para.text = para.text.replace(placeholder, str(value))
#
#             # Save the modified document
#             doc.save(file_path)
#
#             # Convert to PDF if the option is selected
#             if self.pdf_var.get():
#                 pdf_path = os.path.splitext(file_path)[0] + '.pdf'
#                 convert(file_path, pdf_path)
#
#         messagebox.showinfo("Success", "Files generated successfully")
#
# if __name__ == "__main__":
#     app = Replace()
#     app.mainloop()
